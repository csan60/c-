from langchain_community.document_loaders import UnstructuredMarkdownLoader, TextLoader
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter, MarkdownHeaderTextSplitter
from sentence_transformers import SentenceTransformer
from langchain.schema import Document
import os
import fitz  # PyMuPDF 用于处理 PDF 文件
import docx  # python-docx 用于处理 DOCX 文件
import re

# 智能文本分块器
class SmartTextSplitter:
    def __init__(self, chunk_size=1000, chunk_overlap=200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
    def split_text(self, text, file_type="txt"):
        """根据文件类型使用不同的分块策略"""
        if file_type == "md":
            return self._split_markdown(text)
        elif file_type == "pdf":
            return self._split_pdf(text)
        else:
            return self._split_general_text(text)
    
    def _split_markdown(self, text):
        """Markdown文件的智能分块"""
        # 使用MarkdownHeaderTextSplitter按标题分块
        headers_to_split_on = [
            ("#", "标题1"),
            ("##", "标题2"),
            ("###", "标题3"),
            ("####", "标题4"),
        ]
        markdown_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=headers_to_split_on,
            return_each_line=False,
        )
        
        # 先按标题分块
        header_splits = markdown_splitter.split_text(text)
        
        # 如果分块太大，再按段落分块
        final_splits = []
        for split in header_splits:
            if len(split.page_content) > self.chunk_size * 2:
                # 按段落进一步分割
                paragraphs = self._split_by_paragraphs(split.page_content)
                final_splits.extend(paragraphs)
            else:
                final_splits.append(split)
        
        return final_splits
    
    def _split_pdf(self, text):
        """PDF文件的智能分块"""
        # 按段落和句子分块
        paragraphs = self._split_by_paragraphs(text)
        final_splits = []
        
        for para in paragraphs:
            if len(para.page_content) > self.chunk_size:
                # 按句子进一步分割
                sentences = self._split_by_sentences(para.page_content)
                final_splits.extend(sentences)
            else:
                final_splits.append(para)
        
        return final_splits
    
    def _split_general_text(self, text):
        """通用文本的智能分块"""
        # 先按段落分块
        paragraphs = self._split_by_paragraphs(text)
        final_splits = []
        
        for para in paragraphs:
            if len(para.page_content) > self.chunk_size:
                # 按句子进一步分割
                sentences = self._split_by_sentences(para.page_content)
                final_splits.extend(sentences)
            else:
                final_splits.append(para)
        
        return final_splits
    
    def _split_by_paragraphs(self, text):
        """按段落分割文本"""
        paragraphs = re.split(r'\n\s*\n', text.strip())
        docs = []
        current_chunk = ""
        
        for para in paragraphs:
            if len(current_chunk) + len(para) <= self.chunk_size:
                current_chunk += para + "\n\n"
            else:
                if current_chunk:
                    docs.append(Document(page_content=current_chunk.strip()))
                current_chunk = para + "\n\n"
        
        if current_chunk:
            docs.append(Document(page_content=current_chunk.strip()))
        
        return docs
    
    def _split_by_sentences(self, text):
        """按句子分割文本"""
        # 中文句子分割
        sentences = re.split(r'[。！？；]', text)
        docs = []
        current_chunk = ""
        
        for sentence in sentences:
            if len(current_chunk) + len(sentence) <= self.chunk_size:
                current_chunk += sentence + "。"
            else:
                if current_chunk:
                    docs.append(Document(page_content=current_chunk.strip()))
                current_chunk = sentence + "。"
        
        if current_chunk:
            docs.append(Document(page_content=current_chunk.strip()))
        
        return docs

# 加载 DOCX 文件
def load_docx(file_path):
    doc = docx.Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return "\n".join(full_text)

# 加载 PDF 文件
def load_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

# 加载文件
def load_documents(dir_path):
    docs = []
    print(f"正在加载目录: {dir_path}")  # 调试输出
    
    # 初始化智能分块器
    smart_splitter = SmartTextSplitter(chunk_size=600, chunk_overlap=200)
    
    for filename in os.listdir(dir_path):
        file_path = os.path.join(dir_path, filename)
        file_type = filename.split('.')[-1].lower()
        
        try:
            if file_type == 'md':
                print(f"加载 Markdown 文件: {filename}")  # 调试输出
                loader = UnstructuredMarkdownLoader(file_path)
                raw_docs = loader.load()
                # 使用智能分块
                for doc in raw_docs:
                    split_docs = smart_splitter.split_text(doc.page_content, "md")
                    docs.extend(split_docs)
                    
            elif file_type == 'txt':
                print(f"加载文本文件: {filename}")  # 调试输出
                loader = TextLoader(file_path)
                raw_docs = loader.load()
                # 使用智能分块
                for doc in raw_docs:
                    split_docs = smart_splitter.split_text(doc.page_content, "txt")
                    docs.extend(split_docs)
                    
            elif file_type == 'pdf':
                print(f"加载 PDF 文件: {filename}")  # 调试输出
                pdf_text = load_pdf(file_path)
                # 使用智能分块
                split_docs = smart_splitter.split_text(pdf_text, "pdf")
                docs.extend(split_docs)
                
            elif file_type == 'docx':
                print(f"加载 DOCX 文件: {filename}")  # 调试输出
                docx_text = load_docx(file_path)
                # 使用智能分块
                split_docs = smart_splitter.split_text(docx_text, "txt")
                docs.extend(split_docs)
                
        except Exception as e:
            print(f"处理文件 {filename} 时出错: {str(e)}")
            continue
    
    print(f"总共加载了 {len(docs)} 个文档块")  # 调试输出
    return docs

# 生成向量知识库
def create_vector_db(documents, vector_kb_folder):
    if not documents:
        print("没有文档需要处理")
        return None

    try:
        # 使用 m3e-base 模型（safetensors格式，避免PyTorch版本问题）
        embeddings = HuggingFaceEmbeddings(
            model_name='/data-extend/wangqianxu/wqxspace/ITAP/model/m3e-base',
            model_kwargs={'device': 'cpu'},  # 强制使用CPU避免CUDA问题
            encode_kwargs={'normalize_embeddings': True}
        )
        print("✅ 成功加载 m3e-base 模型")
    except Exception as e:
        print(f"加载m3e-base模型失败: {e}")
        print("尝试使用备用方案...")
        # 备用方案：使用SentenceTransformer并手动处理
        try:
            from sentence_transformers import SentenceTransformer
            model = SentenceTransformer('/data-extend/wangqianxu/wqxspace/ITAP/model/m3e-base')
            
            # 手动创建嵌入
            texts = [doc.page_content for doc in documents]
            embeddings_list = model.encode(texts, normalize_embeddings=True)
            
            # 创建自定义嵌入类
            class CustomEmbeddings:
                def __init__(self, model):
                    self.model = model
                
                def embed_documents(self, texts):
                    return self.model.encode(texts, normalize_embeddings=True)
                
                def embed_query(self, text):
                    return self.model.encode([text], normalize_embeddings=True)[0]
            
            embeddings = CustomEmbeddings(model)
            print("✅ 成功加载 m3e-base 模型（备用方案）")
        except Exception as e2:
            print(f"备用方案也失败: {e2}")
            return None

    # 创建或更新 FAISS 向量数据库
    try:
        vectordb = FAISS.from_documents(
            documents=documents,
            embedding=embeddings
        )
        
        # 保存到指定目录
        vectordb.save_local(vector_kb_folder)
        
        print(f"FAISS向量数据库已创建或更新: {vector_kb_folder}")  # 调试输出
        print(f"向量数据库包含 {vectordb.index.ntotal} 个向量")  # 调试输出
        return vectordb
    except Exception as e:
        print(f"创建向量数据库失败: {e}")
        return None

def update_knowledge_db(session_id, isTeacher=False, courseID=None, lessonNum=None):
    """
    更新知识库
    :param session_id: 会话ID
    :param isTeacher: 是否为教师模式
    :param courseID: 课程ID（教师模式下必填）
    :param lessonNum: 课时号（教师模式下必填）
    """
    # 根据isTeacher决定存储路径
    if isTeacher:
        # 教师模式：存储在Teachers目录下的session_id/courseID/lessonNum文件夹中
        if not courseID:
            print("教师模式下courseID不能为空")
            return None
        if not lessonNum:
            print("教师模式下lessonNum不能为空")
            return None
        session_folder = f"/data-extend/wangqianxu/wqxspace/ITAP/base_knowledge/Teachers/{session_id}/{courseID}/{lessonNum}"
    else:
        # 学生模式：存储在Students目录下的session_id文件夹中
        session_folder = f"/data-extend/wangqianxu/wqxspace/ITAP/base_knowledge/Students/{session_id}"
    
    vector_kb_folder = os.path.join(session_folder, "vector_kb")
    os.makedirs(vector_kb_folder, exist_ok=True)

    print(f"开始更新知识库，session_id: {session_id}, isTeacher: {isTeacher}, courseID: {courseID}, lessonNum: {lessonNum}")  # 调试输出

    # 加载文档并生成或更新向量库
    docs = load_documents(session_folder)
    if not docs:
        print("没有找到任何支持的文档类型，请确认上传的文件类型。")  # 调试输出
        return None
        
    vector_db = create_vector_db(docs, vector_kb_folder)
    return vector_db

# 加载已存在的向量数据库
def load_vector_db(session_id, isTeacher=False, courseID=None, lessonNum=None):
    """
    加载已存在的向量数据库
    :param session_id: 会话ID
    :param isTeacher: 是否为教师模式
    :param courseID: 课程ID（教师模式下必填）
    :param lessonNum: 课时号（教师模式下必填）
    """
    # 根据isTeacher决定存储路径
    if isTeacher:
        # 教师模式：从Teachers目录下的session_id/courseID/lessonNum文件夹中加载
        if not courseID:
            print("教师模式下courseID不能为空")
            return None
        if not lessonNum:
            print("教师模式下lessonNum不能为空")
            return None
        vector_kb_folder = f"/data-extend/wangqianxu/wqxspace/ITAP/base_knowledge/Teachers/{session_id}/{courseID}/{lessonNum}/vector_kb"
    else:
        # 学生模式：从Students目录下的session_id文件夹中加载
        vector_kb_folder = f"/data-extend/wangqianxu/wqxspace/ITAP/base_knowledge/Students/{session_id}/vector_kb"
    
    if not os.path.exists(vector_kb_folder):
        print(f"向量数据库不存在: {vector_kb_folder}")
        return None
    
    try:
        # 使用 m3e-base 模型（safetensors格式，避免PyTorch版本问题）
        embeddings = HuggingFaceEmbeddings(
            model_name='/data-extend/wangqianxu/wqxspace/ITAP/model/m3e-base',
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
        vector_db = FAISS.load_local(vector_kb_folder, embeddings)
        print(f"成功加载向量数据库，包含 {vector_db.index.ntotal} 个向量")
        return vector_db
    except Exception as e:
        print(f"加载向量数据库时出错: {str(e)}")
        print("尝试使用备用方案...")
        try:
            # 备用方案：使用SentenceTransformer
            from sentence_transformers import SentenceTransformer
            model = SentenceTransformer('/data-extend/wangqianxu/wqxspace/ITAP/model/m3e-base')
            
            class CustomEmbeddings:
                def __init__(self, model):
                    self.model = model
                
                def embed_documents(self, texts):
                    return self.model.encode(texts, normalize_embeddings=True)
                
                def embed_query(self, text):
                    return self.model.encode([text], normalize_embeddings=True)[0]
            
            embeddings = CustomEmbeddings(model)
            vector_db = FAISS.load_local(vector_kb_folder, embeddings)
            print(f"成功加载向量数据库（备用方案），包含 {vector_db.index.ntotal} 个向量")
            return vector_db
        except Exception as e2:
            print(f"备用方案也失败: {e2}")
            return None